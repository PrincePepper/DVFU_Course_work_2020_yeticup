import sys
from pathlib import Path

import requests
import json
from datetime import datetime

import threading
from threading import Thread

from PyQt5.Qt import QMainWindow, QDialog, QApplication
from PyQt5 import uic, QtWidgets, QtGui
from PyQt5.QtCore import QSize, Qt, QTimer
from PyQt5.QtWidgets import QLineEdit, QTableWidgetItem, QDialog, QDialogButtonBox, QDesktopWidget
from PyQt5.QtGui import QKeyEvent, QPixmap, QFont, QFontDatabase

from docx import Document
from docxtpl import DocxTemplate

authCompetitionName = ''
authLogin = ''
authPassword = ''

API_AUTH = ('yeti', 'yetiyeti')

USERS_API_URL = 'https://yetiapi.herokuapp.com/api/users/'
COMPETITIONS_API_URL = 'https://yetiapi.herokuapp.com/api/competitions'
PLAYERS_API_URL = 'https://yetiapi.herokuapp.com/api/participants'
TEAMS_API_URL = 'https://yetiapi.herokuapp.com/api/teams'

USERS_API_RESPONSE = requests.get(USERS_API_URL, auth = API_AUTH).json()
COMPETITIONS_API_RESPONSE = requests.get(COMPETITIONS_API_URL, auth = API_AUTH).json()
PLAYERS_API_RESPONSE = requests.get(PLAYERS_API_URL, auth = API_AUTH).json()
TEAMS_API_RESPONSE = requests.get(TEAMS_API_URL, auth = API_AUTH).json()

year = []

stream = None
streamDialog = None

mainWindow = None

playersListWindow = None
teamsListWindow = None

addPlayerDialog = None
deletePlayerDialog = None

DEFAULT_FLAGS = Qt.ItemIsEditable | Qt.ItemIsEnabled

def default_item_constructor(content, flags, _ = None):
    item = QTableWidgetItem(content)
    if flags == DEFAULT_FLAGS:
        item.setFlags(Qt.ItemIsEditable)
        item.setFlags(Qt.ItemIsEnabled)
    else:
        item.setFlags(flags)
    return item

def get_result(url, cellRange):
    if url.split('/')[-1] == 'teams':
        response = TEAMS_API_RESPONSE
    if url.split('/')[-1] == 'participants':
        response = PLAYERS_API_RESPONSE

    names = list()
    scores = list()

    scoreKey = 'score'
    nameKey = 'team_name' if url.split('/')[-1] == 'teams' else 'user_id'

    for obj in response:
        if url.split('/')[-1] == 'participants':
            if obj['year'] == year[0]:
                names.append(USERS_API_RESPONSE[obj['user_id'] - 1]['name'])
                scores.append(obj['score'])
        else:
            names.append(obj['team_name'])
            scores.append(obj['score'])

    result = list()
    i = 0
    while i < len(names):
        if cellRange == [0, 0]:
            if scores[i] >= 0:
                result.append(tuple([names[i], scores[i]]))
        else:
            if i >= cellRange[0] and i < cellRange[1]:
                while len(result) < cellRange[1] and i < len(names):
                    if scores[i] >= 0:
                        result.append(tuple([names[i], scores[i]]))
                    i += 1
        i += 1

    return result

def create_table(table, horizontalHeaderLabels, url, item_constructor = default_item_constructor, flags = DEFAULT_FLAGS, cellRange = [0, 0]):
    table.setColumnCount(len(horizontalHeaderLabels))
    table.setHorizontalHeaderLabels(horizontalHeaderLabels)
    header = table.horizontalHeader()
    header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
    header.setSectionResizeMode(1, QtWidgets.QHeaderView.Stretch)
    table.setRowCount(0)

    result = get_result(url, cellRange)

    for i, row in enumerate(result):
        table.setRowCount(table.rowCount() + 1)
        for j, elem in enumerate(row):
            table.setItem(i, j, item_constructor(str(elem), flags, j))
    return result


class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/login.ui', self)
        self.password.setEchoMode(QLineEdit.Password)
        self.pushButton.clicked.connect(lambda: self.auth()) 
        self.pushButton.setStyleSheet('background: rgb(255,220,0);')
        self.setWindowIcon(QtGui.QIcon('logo.png'))
        self.setWindowTitle("Yeti")

        for competition in COMPETITIONS_API_RESPONSE:
            self.competition_name.addItem(competition['name'])

    def auth(self):
        self.login.setStyleSheet("QLineEdit { color: black; background-color: white;}")
        self.password.setStyleSheet("QLineEdit { color: black; background-color: white;}")

        global authCompetitionName
        global authLogin
        global authPassword
        authCompetitionName = self.competition_name.currentText()
        authLogin = self.login.text()
        authPassword = self.password.text()

        global year
        year = [competition['year'] for competition in COMPETITIONS_API_RESPONSE if competition['name'] == authCompetitionName]

        role = [player['role'] for user in USERS_API_RESPONSE if user['mail'] == authLogin for player in PLAYERS_API_RESPONSE if user['id'] == player['user_id'] and player['role'] == 'O']

        password = [user['password'] for user in USERS_API_RESPONSE if user['password'] == authPassword and user['mail'] == authLogin]

        if year and role and password:
            global stream
            global streamDialog

            global mainWindow

            global playersListWindow
            global teamsListWindow

            global addPlayerDialog
            global deletePlayerDialog
            stream = StreamWindow()
            streamDialog = StreamDialog()

            mainWindow = Main()

            playersListWindow = PlayersListWindow()
            teamsListWindow = TeamsListWindow()

            addPlayerDialog = AddPlayerDialog()
            deletePlayerDialog = DeletePlayerDialog()
            mainWindow.showFullScreen()
            self.close()
        elif not role:
            self.login.setStyleSheet("QLineEdit { color: red; background-color: white;}")
        elif not password:
            self.password.setStyleSheet("QLineEdit { color: red; background-color: white;}")


class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/main_window.ui', self)
        self.streamOn = False
        self.show_players_button.clicked.connect(lambda: playersListWindow.showFullScreen())
        self.show_teams_button.clicked.connect(lambda: teamsListWindow.showFullScreen())
        self.show_results_button.clicked.connect(lambda: self.show_results())
        self.gen_report_button.clicked.connect(lambda: self.gen_report())
        self.exit_button.clicked.connect(lambda: self.close_all_windows())
        self.show_players_button.setStyleSheet('background: rgb(255,220,0);')
        self.show_teams_button.setStyleSheet('background: rgb(255,220,0);')
        self.show_results_button.setStyleSheet('background: rgb(255,220,0);')
        self.gen_report_button.setStyleSheet('background: rgb(255,220,0);')
        self.exit_button.setStyleSheet('background: rgb(255,0,0);')

    def show_results(self):
        if QDesktopWidget().screenCount() > 1:
            if not self.streamOn:
                streamDialog.show()
            else:
                stream.close()
                stream.contentTimer.stop()
                self.show_results_button.setText('Транслировать')
            self.streamOn = not self.streamOn

    def gen_report_table(self, doc, n, url):
        result = get_result(url, [0, 0])

        for i in range(len(result)):
            doc.tables[n].add_row()
            doc.tables[n].cell(i + 1, 0).text = str(i + 1)
            doc.tables[n].cell(i + 1, 1).text = str(result[i][1])
            doc.tables[n].cell(i + 1, 2).text = result[i][0]

    def gen_report(self):
        competition_name = authCompetitionName
        competition_date = datetime.now().date()
        competition_address = [competition['address'] for competition in COMPETITIONS_API_RESPONSE if competition['name'] == authCompetitionName]
        manager_name = [user['name'] for user in USERS_API_RESPONSE if user['mail'] == authLogin]

        doc = DocxTemplate('report_template.docx')
        context = { 'competition_name' : competition_name, 'competition_date' : competition_date,
                    'competition_address' : competition_address[0], 'manager_name' : manager_name[0] }
        doc.render(context)
        doc.save('temp_report.docx')

        doc = Document('temp_report.docx')

        self.gen_report_table(doc, 0, TEAMS_API_URL)
        self.gen_report_table(doc, 1, PLAYERS_API_URL)

        doc.save('report.docx')
        Path('temp_report.docx').unlink()

    def close_all_windows(self):
        stream.close()
        playersListWindow.close()
        teamsListWindow.close()
        addPlayerDialog.close()
        deletePlayerDialog.close()
        streamDialog.close()
        mainWindow.close()


class PlayersListWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/players.ui', self)
        self.flags = Qt.ItemIsSelectable | Qt.ItemIsEditable | Qt.ItemIsEnabled
        self.result = create_table(self.table, ['Участник', 'Очки'], PLAYERS_API_URL, self.item_constructor, self.flags)

        self.add_button.clicked.connect(lambda: addPlayerDialog.show())
        self.delete_button.clicked.connect(lambda: deletePlayerDialog.show())
        self.save_button.clicked.connect(lambda: self.save())
        self.to_main_window_button.clicked.connect(lambda: self.close())
        self.add_button.setStyleSheet('background: rgb(255,220,0);')
        self.delete_button.setStyleSheet('background: rgb(255,0,0);')
        self.save_button.setStyleSheet('background: rgb(255,220,0);')
        self.to_main_window_button.setStyleSheet('background: rgb(255,220,0);')

        self.table.keyPressEvent = self.onKeyPress

    def item_constructor(self, content, flags, n):
        item = QTableWidgetItem(content)
        if n != 0:
            item.setFlags(flags)
        else:
            item.setFlags(Qt.ItemIsEditable)
            item.setFlags(Qt.ItemIsEnabled)
        return item

    def onKeyPress(self, key: QKeyEvent):
        if key.key() == Qt.Key_Return:
            self.save()

    def save(self):
        newTable = []
        oldTable = self.result
        for i in range(len(oldTable)):
            if self.table.item(i, 0) is None:
                continue
            name = self.table.item(i, 0).text()
            score = self.table.item(i, 1).text()
            newTable.append((name, score))
        self.result = newTable
        for i, row in enumerate(newTable):
            if row[1] != oldTable[i][1]:
                player = 0
                for bplayer in PLAYERS_API_RESPONSE:
                    for user in USERS_API_RESPONSE:
                        if user['name'] == row[0] and bplayer['score'] == oldTable[i][1] and bplayer['year'] == year[0]:
                            player = bplayer
                player['score'] = row[1]
                package = Thread(target = self.send_data, args = (player, ))
                package.start()

    def send_data(self, data):
        response = requests.put(PLAYERS_API_URL + '/' + str(data['id']) + '/', data, auth = API_AUTH)


class TeamsListWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/teams.ui', self)
        create_table(self.table, ['Название', 'Очки'], TEAMS_API_URL)

        self.to_main_window_button.clicked.connect(lambda: self.close())
        self.to_main_window_button.setStyleSheet('background: rgb(255,220,0);')


class AddPlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/add_player.ui', self)
        self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

    def accept(self):
        data = { "name" : self.name.toPlainText(), "password" : self.password.toPlainText(), "mail" : self.email.toPlainText(),
                 "address" : self.address.toPlainText(), "phone" : self.phone.toPlainText(), "photo" : None }
        package = Thread(target = self.send_data, args = (data, ))
        package.start()
        self.close()

    def send_data(self, data):
        response = requests.post(USERS_API_URL, data, auth = API_AUTH)


class DeletePlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/delete_player.ui', self)
        self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

    def accept(self):
        rows = playersListWindow.table.selectionModel().selectedIndexes()
        for row in rows:
            playersListWindow.table.removeRow(rows[0].row())
            player_id = 0
            for player in PLAYERS_API_RESPONSE:
                for user in USERS_API_RESPONSE:
                    if user['name'] == playersListWindow.result[row.row()][0]:
                        if player['score'] == playersListWindow.result[row.row()][1]:
                            if player['year'] == year[0]:
                                player_id = player['id']

            package = Thread(target = self.send_data, args = (PLAYERS_API_URL + '/' + str(player_id), ))
            package.start()
        self.close()

    def send_data(self, url):
        response = requests.delete(url, auth = API_AUTH)


class StreamDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/stream_dialog.ui', self)
        self.screens = [self.screen1, self.screen2, self.screen3, self.screen4, self.screen5, self.screen6]
        k = 0
        for i, screenButton in enumerate(self.screens):
            if i > QDesktopWidget().screenCount() - 2:
                screenButton.hide()
            else:
                k = i
                screenButton.clicked.connect(lambda: self.showStream(k + 1))
                screenButton.setStyleSheet('background: rgb(255,220,0);')
        self.cancel.clicked.connect(self.close)
        self.cancel.setStyleSheet('background: rgb(255,220,0);')

    def showStream(self, screenNumber):
        stream.show_content()
        monitor = QDesktopWidget().screenGeometry(screenNumber)
        stream.move(monitor.left(), monitor.top())
        stream.showFullScreen()
        mainWindow.show_results_button.setText('Завершить трансляцию')
        self.close()


class StreamWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/for_stream.ui', self)
        self.label.setPixmap(QPixmap('logo.png'))
        self.label.setAlignment(Qt.AlignCenter)
        self.contents = [self.show_image, self.show_players, self.show_teams]
        self.currentContent = 0
        self.contentTimer = QTimer(self, timeout = lambda: self.contents[self.currentContent]())
        self.step = 17
        self.playersCellRange = [0, self.step]
        self.teamsCellRange = [0, self.step]

    def show_content(self):
        self.show_image()
        self.contentTimer.start(4000)

    def show_image(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.results_table.hide()
        self.label.show()

    def show_players(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.label.hide()
        self.results_table.clear()
        create_table(self.results_table, ['Участник', 'Очки'], PLAYERS_API_URL, cellRange = self.playersCellRange)
        if self.results_table.rowCount() == self.playersCellRange[1] - self.playersCellRange[0]:
            self.playersCellRange[0] += self.step
            self.playersCellRange[1] += self.step
        else:
            self.playersCellRange = [0, self.step]
        self.results_table.show()

    def show_teams(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.label.hide()
        self.results_table.clear()
        create_table(self.results_table, ['Название', 'Очки'], TEAMS_API_URL, cellRange = self.teamsCellRange)
        if self.results_table.rowCount() == self.teamsCellRange[1] - self.teamsCellRange[0]:
            self.teamsCellRange[0] += self.step
            self.teamsCellRange[1] += self.step
        else:
            self.teamsCellRange = [0, self.step]
        self.results_table.show()


app = QApplication(sys.argv)
app.setFont(QFont("Roboto", 11))
loginWindow = LoginWindow()
loginWindow.show()

sys.exit(app.exec_())

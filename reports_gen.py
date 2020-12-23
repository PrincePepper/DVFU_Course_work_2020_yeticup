from docx import Document
import sqlite3

doc = Document("report_template.docx")
doc.tables

data = sqlite3.connect('players.db')
command = data.cursor()
players = command.execute("SELECT name, score FROM players;").fetchall()

for i in range(len(players)):
    doc.tables[1].add_row()
    doc.tables[1].cell(i + 1, 0).text = str(i + 1)
    doc.tables[1].cell(i + 1, 1).text = str(players[i][1])
    doc.tables[1].cell(i + 1, 2).text = players[i][0]

data = sqlite3.connect('teams.db')
command = data.cursor()
teams = command.execute("SELECT name, score FROM teams;").fetchall()
for i in range(len(teams)):
    doc.tables[0].add_row()
    doc.tables[0].cell(i + 1, 0).text = str(i + 1)
    doc.tables[0].cell(i + 1, 1).text = teams[i][0]
    doc.tables[0].cell(i + 1, 2).text = str(teams[i][1])

doc.save("report.docx")
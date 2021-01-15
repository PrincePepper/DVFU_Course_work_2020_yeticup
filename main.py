import sys
from pathlib import Path

import requests
import json
from datetime import datetime

import threading
from threading import Thread

from PyQt5.Qt import QMainWindow, QDialog, QApplication
from PyQt5 import uic, QtWidgets
from PyQt5.QtCore import QSize, Qt, QTimer
from PyQt5.QtWidgets import QLineEdit, QTableWidgetItem, QDialog, QDialogButtonBox, QDesktopWidget
from PyQt5.QtGui import QKeyEvent, QPixmap

from docx import Document
from docxtpl import DocxTemplate

authCompetitionName = ''
authLogin = ''
authPassword = ''

USERS_API_URL = 'https://yetiapi.herokuapp.com/api/users/'
COMPETITIONS_API_URL = 'https://yetiapi.herokuapp.com/api/competitions'
PLAYERS_API_URL = 'https://yetiapi.herokuapp.com/api/participants'
TEAMS_API_URL = 'https://yetiapi.herokuapp.com/api/teams'

USERS_API_RESPONSE = requests.get(USERS_API_URL).json()
COMPETITIONS_API_RESPONSE = requests.get(COMPETITIONS_API_URL).json()
PLAYERS_API_RESPONSE = requests.get(PLAYERS_API_URL).json()
TEAMS_API_RESPONSE = requests.get(TEAMS_API_URL).json()

year = [2000]

DEFAULT_FLAGS = Qt.ItemIsEditable | Qt.ItemIsEnabled

def default_item_constructor(content, flags, _ = None):
    item = QTableWidgetItem(content)
    if flags == DEFAULT_FLAGS:
        item.setFlags(Qt.ItemIsEditable)
        item.setFlags(Qt.ItemIsEnabled)
    else:
        item.setFlags(flags)
    return item

def get_result(url, cellRange):
    if url.split('/')[-1] == 'teams':
        response = TEAMS_API_RESPONSE
    if url.split('/')[-1] == 'participants':
        response = PLAYERS_API_RESPONSE

    names = list()
    scores = list()

    scoreKey = 'score'
    nameKey = 'team_name' if url.split('/')[-1] == 'teams' else 'user_id'

    for obj in response:
        if url.split('/')[-1] == 'participants':
            if obj['year'] == year[0]:
                names.append(USERS_API_RESPONSE[obj['user_id'] - 1]['name'])
                scores.append(obj['score'])
        else:
            names.append(obj['team_name'])
            scores.append(obj['score'])

    result = list()
    i = 0
    while i < len(names):
        if cellRange == [0, 0]:
            if scores[i] >= 0:
                result.append(tuple([names[i], scores[i]]))
        else:
            if i >= cellRange[0] and i < cellRange[1]:
                while len(result) < cellRange[1] and i < len(names):
                    if scores[i] >= 0:
                        result.append(tuple([names[i], scores[i]]))
                    i += 1
        i += 1

    return result

def create_table(table, horizontalHeaderLabels, url, item_constructor = default_item_constructor, flags = DEFAULT_FLAGS, cellRange = [0, 0]):
    table.setColumnCount(len(horizontalHeaderLabels))
    table.setHorizontalHeaderLabels(horizontalHeaderLabels)
    header = table.horizontalHeader()
    header.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeToContents)
    header.setSectionResizeMode(1, QtWidgets.QHeaderView.Stretch)
    table.setRowCount(0)

    result = get_result(url, cellRange)

    for i, row in enumerate(result):
        table.setRowCount(table.rowCount() + 1)
        for j, elem in enumerate(row):
            table.setItem(i, j, item_constructor(str(elem), flags, j))
    return result


class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/login.ui', self)
        self.password.setEchoMode(QLineEdit.Password)
        self.pushButton.clicked.connect(lambda: self.auth())

    def auth(self):
        self.competition_name.setStyleSheet("QLineEdit { color: black; background-color: white;}")
        self.login.setStyleSheet("QLineEdit { color: black; background-color: white;}")
        self.password.setStyleSheet("QLineEdit { color: black; background-color: white;}")

        global authCompetitionName
        global authLogin
        global authPassword
        authCompetitionName = self.competition_name.text()
        authLogin = self.login.text()
        authPassword = self.password.text()

        global year
        year = [competition['year'] for competition in COMPETITIONS_API_RESPONSE if competition['name'] == authCompetitionName]

        role = [player['role'] for user in USERS_API_RESPONSE if user['login'] == authLogin for player in PLAYERS_API_RESPONSE if user['id'] == player['user_id'] and player['role'] == 'O']

        password = [user['password'] for user in USERS_API_RESPONSE if user['password'] == authPassword and user['login'] == authLogin]

        if year and role and password:
            mainWindow.showFullScreen()
            self.close()
        elif not year:
            self.competition_name.setStyleSheet("QLineEdit { color: red; background-color: white;}")
        elif not role:
            self.login.setStyleSheet("QLineEdit { color: red; background-color: white;}")
        elif not password:
            self.password.setStyleSheet("QLineEdit { color: red; background-color: white;}")


class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/main_window.ui', self)
        self.streamOn = False
        self.show_players_button.clicked.connect(lambda: playersListWindow.showFullScreen())
        self.show_teams_button.clicked.connect(lambda: teamsListWindow.showFullScreen())
        self.show_results_button.clicked.connect(lambda: self.show_results())
        self.gen_report_button.clicked.connect(lambda: self.gen_report())
        self.exit_button.clicked.connect(lambda: self.close_all_windows())

    def show_results(self):
        if not self.streamOn:
            stream.show_content()
            monitor = QDesktopWidget().screenGeometry(1)
            stream.move(monitor.left(), monitor.top())
            stream.showFullScreen()
            self.show_results_button.setText('Завершить трансляцию')
        else:
            stream.close()
            stream.contentTimer.stop()
            self.show_results_button.setText('Транслировать')
        self.streamOn = not self.streamOn

    def gen_report_table(self, doc, n, url):
        result = get_result(url)

        for i in range(len(result)):
            doc.tables[n].add_row()
            doc.tables[n].cell(i + 1, 0).text = str(i + 1)
            doc.tables[n].cell(i + 1, 1).text = str(result[i][1])
            doc.tables[n].cell(i + 1, 2).text = result[i][0]

    def gen_report(self):
        competition_name = authCompetitionName
        competition_date = datetime.now().date()
        competition_address = [competition['address'] for competition in COMPETITIONS_API_RESPONSE if competition['name'] == authCompetitionName]
        manager_name = [user['name'] for user in USERS_API_RESPONSE if user['login'] == authLogin]

        doc = DocxTemplate('report_template.docx')
        context = { 'competition_name' : competition_name, 'competition_date' : competition_date,
                    'competition_address' : competition_address[0], 'manager_name' : manager_name[0] }
        doc.render(context)
        doc.save('temp_report.docx')

        doc = Document('temp_report.docx')

        self.gen_report_table(doc, 0, TEAMS_API_URL)
        self.gen_report_table(doc, 1, PLAYERS_API_URL)

        doc.save('report.docx')
        Path('temp_report.docx').unlink()

    def close_all_windows(self):
        stream.close()
        playersListWindow.close()
        teamsListWindow.close()
        addPlayerDialog.close()
        deletePlayerDialog.close()
        # duplicatePlayerDialog.close()
        mainWindow.close()


class PlayersListWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/players.ui', self)
        self.flags = Qt.ItemIsSelectable | Qt.ItemIsEditable | Qt.ItemIsEnabled
        self.result = create_table(self.table, ['Участник', 'Очки'], PLAYERS_API_URL, self.item_constructor, self.flags)

        self.add_button.clicked.connect(lambda: addPlayerDialog.show())
        self.delete_button.clicked.connect(lambda: deletePlayerDialog.show())
        # self.duplicate_button.clicked.connect(lambda: duplicatePlayerDialog.show())

        self.save_button.clicked.connect(lambda: self.save())

        self.to_main_window_button.clicked.connect(lambda: self.close())

        self.table.keyPressEvent = self.onKeyPress

    def item_constructor(self, content, flags, n):
        item = QTableWidgetItem(content)
        if n != 0:
            item.setFlags(flags)
        else:
            item.setFlags(Qt.ItemIsEditable)
            item.setFlags(Qt.ItemIsEnabled)
        return item

    def onKeyPress(self, key: QKeyEvent):
        if key.key() == Qt.Key_Return:
            self.save()

    def save(self):
        newTable = []
        oldTable = self.result
        for i in range(len(oldTable)):
            name = self.table.item(i, 0).text()
            score = self.table.item(i, 1).text()
            newTable.append((name, score))
        self.result = newTable
        for i, row in enumerate(newTable):
            if row[1] != oldTable[i][1]:
                player = 0
                for bplayer in PLAYERS_API_RESPONSE:
                    for user in USERS_API_RESPONSE:
                        if user['name'] == row[0] and bplayer['score'] == oldTable[i][1] and bplayer['year'] == year[0]:
                            player = bplayer
                player['score'] = row[1]
                package = Thread(target = self.send_data, args = (player, ))
                package.start()

    def send_data(self, data):
        response = requests.put(PLAYERS_API_URL + '/' + str(data['id']) + '/', data)


class TeamsListWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/teams.ui', self)
        create_table(self.table, ['Название', 'Очки'], TEAMS_API_URL)

        self.to_main_window_button.clicked.connect(lambda: self.close())


class AddPlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/add_player.ui', self)
        self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

    def accept(self):
        data = { "name" : self.name.toPlainText(), "login" : self.login.toPlainText(), "password" : self.password.toPlainText(), "mail" : self.email.toPlainText(),
                 "address" : self.address.toPlainText(), "phone" : self.phone.toPlainText(), "photo" : None }
        package = Thread(target = self.send_data, args = (data, ))
        package.start()
        self.close()
    
    def send_data(self, data):
        response = requests.post(USERS_API_URL, data)


class DeletePlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/delete_player.ui', self)
        self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

    def accept(self):
        rows = playersListWindow.table.selectionModel().selectedIndexes()
        for row in rows:
            playersListWindow.table.removeRow(rows[0].row())
            player_id = 0
            for player in PLAYERS_API_RESPONSE:
                for user in USERS_API_RESPONSE:
                    if user['name'] == playersListWindow.result[row.row()][0]:
                        if player['score'] == playersListWindow.result[row.row()][1]:
                            if player['year'] == year[0]:
                                player_id = player['id']

            package = Thread(target = self.send_data, args = (PLAYERS_API_URL + '/' + str(player_id), ))
            package.start()
        self.close()

    def send_data(self, url):
        response = requests.delete(url)

# class DuplicatePlayerDialog(QDialog):
#     def __init__(self):
#         super().__init__()
#         uic.loadUi('ui_files/duplication.ui', self)
#         self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
#         self.buttonBox.accepted.connect(self.accept)
#         self.buttonBox.rejected.connect(self.reject)

#     def accept(self):
#         self.old_login.setStyleSheet("QLineEdit { color: black; background-color: white;}")
#         old_user = [user for user in USERS_API_RESPONSE if user['login'] == self.old_login.text()]
#         if old_user:
#             data = { "name" : old_user[0]['name'], "login" : self.new_login.text(), "password" : self.new_password.text(), "mail" : self.new_mail.text(),
#                     "address" : old_user[0]['address'], "phone" : old_user[0]['phone'], "photo" : old_user[0]['photo'] }
#             package = Thread(target = self.send_data, args = (data, ))
#             package.start()
#         else:
#             self.old_login.setStyleSheet("QLineEdit { color: red; background-color: white;}")
#         self.close()

#     def send_data(self, data):
#         response = requests.post(USERS_API_URL, data)

class StreamWindow(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('ui_files/for_stream.ui', self)
        self.label.setPixmap(QPixmap('logo.png'))
        self.label.setAlignment(Qt.AlignCenter)
        self.contents = [self.show_image, self.show_players, self.show_teams]
        self.currentContent = 0
        self.contentTimer = QTimer(self, timeout = lambda: self.contents[self.currentContent]())
        self.playersCellRange = [0, 19]
        self.teamsCellRange = [0, 19]
        self.step = 19

    def show_content(self):
        self.show_image()
        self.contentTimer.start(4000)

    def show_image(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.results_table.hide()
        self.label.show()

    def show_players(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.label.hide()
        self.results_table.clear()
        create_table(self.results_table, ['Участник', 'Очки'], PLAYERS_API_URL, cellRange = self.playersCellRange)
        if self.results_table.rowCount() == self.playersCellRange[1] - self.playersCellRange[0]:
            self.playersCellRange[0] += self.step
            self.playersCellRange[1] += self.step
        else:
            self.playersCellRange = [0, 19]
        self.results_table.show()

    def show_teams(self):
        self.currentContent = (self.currentContent + 1) % len(self.contents)
        self.label.hide()
        self.results_table.clear()
        create_table(self.results_table, ['Название', 'Очки'], TEAMS_API_URL, cellRange = self.teamsCellRange)
        if self.results_table.rowCount() == self.teamsCellRange[1] - self.teamsCellRange[0]:
            self.teamsCellRange[0] += self.step
            self.teamsCellRange[1] += self.step
        else:
            self.teamsCellRange = [0, 19]
        self.results_table.show()


app = QApplication(sys.argv)
loginWindow = LoginWindow()
loginWindow.show()

stream = StreamWindow()

mainWindow = Main()

playersListWindow = PlayersListWindow()
teamsListWindow = TeamsListWindow()

addPlayerDialog = AddPlayerDialog()
deletePlayerDialog = DeletePlayerDialog()
# duplicatePlayerDialog = DuplicatePlayerDialog()
sys.exit(app.exec_())
